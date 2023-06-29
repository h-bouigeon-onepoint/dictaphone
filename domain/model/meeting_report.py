import os
from docx import Document
from flask import current_app as app
from werkzeug.utils import secure_filename
from services.constants import Constants
from services.handle_ppt import PPTWriter
from services.read_bytes_files import read_bytes_file
from services.handle_json import read_json_file, write_json_file
from services.handle_pickle import read_pickle_file, write_pickle_file
from services.api.text_to_meeting_report import text_to_meeting_report
from domain.model.meeting_report_data import MeetingReportData


class MeetingReport:
    def __init__(
        self,
        filename: str,
        transcript: str = "",
        language: str = "french",
        meeting_report_json: dict = Constants.MODEL_JSON_MEETING_REPORT,
    ) -> None:
        self.filename = secure_filename(str(filename))
        self.filepath = os.path.join(app.config["MEETING_REPORT_FOLDER"], self.filename)
        self.metadata_path = os.path.join(app.config["METADATA_FOLDER"], self.filename)
        self.language = language
        self.transcript = transcript
        self.meeting_report_json = MeetingReportData(**meeting_report_json)
        self.content = ""
        self.ppt_path = os.path.join(app.config["PPT_FOLDER"], self.filename) + ".pptx"

    def get_meeting_report_from_transcript(self) -> None:
        print("LANGUAGE :", self.language)
        self.meeting_report_json = MeetingReportData(
            **text_to_meeting_report(self.transcript, self.language)
        )

    def write_to_json(self) -> None:
        write_json_file(f"{self.filepath}.json", self.meeting_report_json.to_dict())

    def read_from_json(self) -> None:
        self.meeting_report_json = MeetingReportData(
            **read_json_file(f"{self.filepath}.json")
        )

    def write_metadata_to_pickle(self) -> None:
        write_pickle_file(f"{self.metadata_path}.pickle", {"language": self.language})

    def read_metadata_from_pickle(self) -> None:
        metadata = read_pickle_file(f"{self.metadata_path}.pickle")
        self.language = metadata["language"]

    def write_to_ppt(self) -> None:
        meeting_report_power_point = PPTWriter(self.filename, self.meeting_report_json)
        meeting_report_power_point.write_to_ppt(self.ppt_path)

    def get_ppt_file(self) -> bytes:
        return read_bytes_file(self.ppt_path)

    def write_to_word(self) -> None:
        # TO DO: nicer word export
        self.word_path = (
            os.path.join(app.config["WORD_FOLDER"], self.filename) + ".docx"
        )
        document = Document()
        document.add_paragraph("Titre : " + self.meeting_report_json.title_meeting)
        document.add_paragraph("Participants :")
        for participant in self.meeting_report_json.participants:
            document.add_paragraph(participant)
        document.add_paragraph("Ordre du jour :")
        for agenda in self.meeting_report_json.agenda:
            document.add_paragraph(agenda)
        document.add_paragraph("Éléments discutés :")
        for element in self.meeting_report_json.elements_discussed:
            document.add_paragraph(element)
        document.add_paragraph("Actions :")
        for action in self.meeting_report_json.action_to_be_taken:
            document.add_paragraph(action["action"])
            document.add_paragraph(action["responsable"])
            document.add_paragraph(action["deadline"])
        document.add_paragraph("Notes:\n" + self.meeting_report_json.notes)
        document.save(self.word_path)

    def get_word_file(self) -> bytes:
        return read_bytes_file(self.word_path)
